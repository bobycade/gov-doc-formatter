# -*- coding: utf-8 -*-
"""Chinese Government Document Formatter (GB/T 9704-2012).

Converts markdown to .docx with proper 公文 formatting.
Supports: multi-level headings, tables, blockquotes, lists, inline bold,
mixed Chinese/English fonts.

Usage:
    python format_doc.py --input source.md --output result.docx
    python format_doc.py --input source.md          # auto name
    python format_doc.py --json data.json            # JSON mode (legacy)
"""

import re
import argparse
import json
import os
import shutil
import tempfile
import importlib.util
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Font constants ──────────────────────────────────────────────
FZ = '方正小标宋简体'   # main title
HT = '黑体'             # chapter headings, table headers
KT = '楷体_GB2312'      # section headings, notes
FS = '仿宋_GB2312'      # body text
EN = 'Times New Roman'  # numbers / English

# ── Helpers ─────────────────────────────────────────────────────

def _set_outline(p, level):
    """Set outline level on a paragraph via direct XML (python-docx property is unreliable)."""
    pPr = p._element.get_or_add_pPr()
    ol = parse_xml(f'<w:outlineLvl {nsdecls("w")} w:val="{level}"/>')
    pPr.append(ol)


def _set_widow_control(p):
    """Disable widow/orphan control on a paragraph (match reference docs)."""
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:widowControl {nsdecls("w")} w:val="0"/>'))


def _rf(run, name=FS, east=None, size=Pt(16), bold=False):
    """Apply font properties to a run."""
    if east is None:
        east = name
    run.font.name = name
    run.font.size = size
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east)


def _add_mixed(p, text, name, east, size, bold=False):
    """Add text to paragraph, auto-switching EN font for numbers/English."""
    parts = re.split(r'([a-zA-Z0-9%\-\.㎡‰＋≥≤～/°℃]+)', text)
    for part in parts:
        if not part:
            continue
        r = p.add_run(part)
        if re.match(r'^[a-zA-Z0-9%\-\.㎡‰＋≥≤～/°℃]+$', part):
            _rf(r, EN, EN, size, bold)
        else:
            _rf(r, name, east, size, bold)


def _shade(cell, color):
    """Set cell background color."""
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))


def _convert_quotes(text):
    """Convert straight ASCII quotes to Chinese quotation marks ("" → "")."""
    result = []
    dq = False  # double-quote state
    sq = False  # single-quote state
    for ch in text:
        if ch == '"':
            result.append('”' if dq else '“')
            dq = not dq
        elif ch == "'":
            result.append('’' if sq else '‘')
            sq = not sq
        else:
            result.append(ch)
    return ''.join(result)


def _add_rich_text(p, text, name=FS, east=None, size=Pt(16), bold_default=False):
    """Add text with **bold** processing, quote conversion, EN font switching."""
    if east is None:
        east = name
    text = _convert_quotes(text)
    segs = re.split(r'(\*\*.*?\*\*)', text)
    for i, seg in enumerate(segs):
        if not seg:
            continue
        if seg.startswith('**') and seg.endswith('**'):
            _add_mixed(p, seg[2:-2], name, east, size, True)
        else:
            # Strip leading space that follows a **bold** segment (Chinese
            # typesetting does not use whitespace between runs).
            prev_bold = i > 0 and segs[i-1].startswith('**') and segs[i-1].endswith('**')
            if prev_bold:
                seg = seg.lstrip(' ')
            _add_mixed(p, seg, name, east, size, bold_default)


# ── Numbering normalization ──────────────────────────────────────

_CHINESE_DIGITS = '零一二三四五六七八九'


def _to_chinese_num(n):
    """Convert integer 1-99 to Chinese numeral: 1→一, 10→十, 11→十一."""
    if n <= 0:
        return str(n)
    if n <= 9:
        return _CHINESE_DIGITS[n]
    if n == 10:
        return '十'
    if n < 20:
        return '十' + _CHINESE_DIGITS[n % 10]
    if n % 10 == 0:
        return _CHINESE_DIGITS[n // 10] + '十'
    return _CHINESE_DIGITS[n // 10] + '十' + _CHINESE_DIGITS[n % 10]


def _circled_num(n):
    """Convert 1-20 to circled number: 1→①, 20→⑳. Falls back to plain number."""
    if 1 <= n <= 20:
        return chr(0x245F + n)
    return str(n)


def _strip_heading_num(text):
    """Strip leading numbering from heading text.

    Handles: 1.1, 1.1.1, 4.2.1, 1., 1、, 一、, （一）, (1), 1), ①, etc.
    Returns cleaned text without the numbering prefix.
    """
    patterns = [
        r'^\d+(\.\d+)+[\s\.、．]+',      # 1.1, 1.1.1, 4.2.1
        r'^\d+[\s\.、．\)）]+',            # 1. , 1、, 1), 1）
        r'^[一二三四五六七八九十]+[、．]\s*',  # 一、二、
        r'^（[一二三四五六七八九十]+）\s*',    # （一）（二）
        r'^\(?\d+[\)）]\s*',              # (1), 1)
        r'^[①②③④⑤⑥⑦⑧⑨⑩]\s*',          # ①②③
    ]
    for pat in patterns:
        m = re.match(pat, text)
        if m:
            return text[m.end():]
    return text


def _body_core(p, text, indent=True):
    """Common logic for body paragraphs with inline bold."""
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(32) if indent else Pt(0)
    p.paragraph_format.line_spacing = Pt(29)
    _set_widow_control(p)
    _add_rich_text(p, text, FS, FS, Pt(16), False)


# ── Paragraph builders ──────────────────────────────────────────

def add_title(doc, text):
    """Main title: 方正小标宋简体 22pt, centered (not bold per GB/T 9704-2012)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    _add_rich_text(p, text, FZ, FZ, Pt(22), False)


def add_subtitle(doc, text):
    """Subtitle / version line: 楷体 16pt, centered. Supports **bold**."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    _set_widow_control(p)
    _add_rich_text(p, text, KT, KT, Pt(16), False)


def add_h1(doc, text):
    """Chapter heading (一、二、): 黑体 16pt bold."""
    p = doc.add_paragraph()
    _set_outline(p, 0)
    p.paragraph_format.first_line_indent = Pt(32)
    p.paragraph_format.line_spacing = Pt(29)
    _set_widow_control(p)
    _add_rich_text(p, text, HT, HT, Pt(16), True)


def add_h2(doc, text):
    """Section heading (1.1, 2.1): 楷体_GB2312 16pt bold."""
    p = doc.add_paragraph()
    _set_outline(p, 1)
    p.paragraph_format.first_line_indent = Pt(32)
    p.paragraph_format.line_spacing = Pt(29)
    _set_widow_control(p)
    _add_rich_text(p, text, KT, KT, Pt(16), True)


def add_h3(doc, text):
    """Sub-section heading (4.2.1): 仿宋 16pt bold."""
    p = doc.add_paragraph()
    _set_outline(p, 2)
    p.paragraph_format.first_line_indent = Pt(32)
    p.paragraph_format.line_spacing = Pt(29)
    _set_widow_control(p)
    _add_rich_text(p, text, FS, FS, Pt(16), True)


def add_body(doc, text, indent=True):
    """Body: 仿宋 16pt, first-line indent 32pt. Supports **bold**."""
    p = doc.add_paragraph()
    _body_core(p, text, indent)


def add_note(doc, text):
    """Note / blockquote: 楷体 14pt."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(32)
    p.paragraph_format.line_spacing = Pt(29)
    _set_widow_control(p)
    _add_rich_text(p, text, KT, KT, Pt(14), False)


def add_signature(doc, lines):
    """Document ending (落款): issuing authority + date, right-aligned.

    Args:
        lines: list of strings, each right-aligned with appropriate spacing.
    """
    # Two blank lines before signature block
    for _ in range(2):
        sp = doc.add_paragraph()
        sp.paragraph_format.first_line_indent = Pt(0)
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after = Pt(0)
        _set_widow_control(sp)

    for i, text in enumerate(lines):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _set_widow_control(p)
        # Last line (typically the date) gets extra right indent for visual balance
        if i == len(lines) - 1:
            p.paragraph_format.right_indent = Pt(16)
        _add_rich_text(p, text, FS, FS, Pt(16), False)


def add_bullet(doc, text):
    """Bullet item: 仿宋 16pt, prefixed with em-dash."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(32)
    p.paragraph_format.line_spacing = Pt(29)
    _set_widow_control(p)
    _add_rich_text(p, '— ' + text, FS, FS, Pt(16), False)


def add_table(doc, headers, rows, left_content=False):
    """Table: header 黑体 12pt bold (blue bg), body 仿宋 12pt.

    Args:
        left_content: if True, columns beyond the first use left-alignment
                      (for content-heavy tables).
    """
    nc = len(headers)
    tbl = doc.add_table(rows=len(rows) + 1, cols=nc)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    # Repeat header row across pages
    trPr = tbl.rows[0]._tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = parse_xml(f'<w:trPr {nsdecls("w")}/>')
        tbl.rows[0]._tr.insert(0, trPr)
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

    # Header row: 黑体 小四(12pt) bold, blue bg, compact spacing
    for j, hh in enumerate(headers):
        c = tbl.rows[0].cells[j]; c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = Pt(20)
        _set_widow_control(p)
        _add_rich_text(p, hh, HT, HT, Pt(12), True)
        _shade(c, 'D9E2F3')

    # Data rows: 仿宋 小四(12pt), compact spacing
    for i, rd in enumerate(rows):
        for j, v in enumerate(rd):
            c = tbl.rows[i + 1].cells[j]; c.text = ''
            p = c.paragraphs[0]
            if left_content and j > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = Pt(20)
            _set_widow_control(p)
            _add_rich_text(p, str(v), FS, FS, Pt(12), False)

    # Minimal spacer after table
    bp = doc.add_paragraph()
    bp.paragraph_format.first_line_indent = Pt(0)
    bp.paragraph_format.line_spacing = Pt(2)


def add_page_numbers(doc):
    """Add '— PAGE —' footer to all sections."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        r1 = fp.add_run('— '); r1.font.name = EN; r1.font.size = Pt(12)

        from docx.oxml import OxmlElement
        fc1 = OxmlElement('w:fldChar'); fc1.set(qn('w:fldCharType'), 'begin')
        it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve')
        it.text = 'PAGE'
        fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'), 'separate')
        fc3 = OxmlElement('w:fldChar'); fc3.set(qn('w:fldCharType'), 'end')

        rp = fp.add_run(); rp.font.name = EN; rp.font.size = Pt(12)
        rp._r.append(fc1); rp._r.append(it); rp._r.append(fc2); rp._r.append(fc3)

        r2 = fp.add_run(' —'); r2.font.name = EN; r2.font.size = Pt(12)


# ── Markdown parser ─────────────────────────────────────────────

def _parse_table(lines, start):
    """Parse markdown table starting at lines[start]. Returns (headers, rows, next_idx)."""
    raw = []
    i = start
    while i < len(lines) and '|' in lines[i] and lines[i].strip():
        raw.append(lines[i].strip())
        i += 1
    if len(raw) < 2:
        return None, None, i

    headers = [c.strip() for c in raw[0].split('|')[1:-1]]
    data_start = 2 if re.match(r'^[\|\s\-:]+$', raw[1]) else 1
    rows = [[c.strip() for c in rl.split('|')[1:-1]] for rl in raw[data_start:]]
    return headers, rows, i


def _is_content_table(rows):
    """Heuristic: if any non-first column has cells > 20 chars, treat as content table."""
    if not rows:
        return False
    for r in rows[:3]:
        for c in r[1:]:
            if len(c) > 20:
                return True
    return False


def convert_markdown(md_path, docx_path):
    """Convert a markdown file to government-format .docx."""
    if not os.path.exists(md_path):
        raise FileNotFoundError(f"Markdown file not found: {md_path}")

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # Page setup (A4, standard margins)
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(3.7)
    sec.bottom_margin = Cm(3.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.6)

    # Default style
    ns = doc.styles['Normal']
    ns.font.name = FS
    ns.font.size = Pt(16)
    ns.element.rPr.rFonts.set(qn('w:eastAsia'), FS)
    ns.paragraph_format.line_spacing = Pt(29)
    ns.paragraph_format.space_before = Pt(0)
    ns.paragraph_format.space_after = Pt(0)
    ns.paragraph_format.first_line_indent = Pt(32)
    # Widow/orphan control — disabled to match reference docs
    ns._element.get_or_add_pPr().append(parse_xml(f'<w:widowControl {nsdecls("w")} w:val="0"/>'))

    i = 0
    # GB/T 9704 numbering state
    h2_count = 0       # （一）（二）counter, resets per chapter (##)
    h3_count = 0       # 1. 2. 3. counter, resets per section (###)
    heading_depth = 0  # 1=##, 2=###, 3=####, 0=before first heading
    bullet_num = 0     # current - item counter
    plus_num = 0       # current + item counter (Level 5: ①②③)
    body_label_num = 0 # counter for bold body labels, resets per heading
    in_bold_section = False  # True when inside a bold body label section
    while i < len(lines):
        s = lines[i].rstrip('\n').strip()
        if not s:
            i += 1
            continue

        # Reset item counters on lines that aren't list items or headings
        if not s.startswith('- ') and not s.startswith('+ ') and not s.startswith('#'):
            bullet_num = 0
        if not s.startswith('+ ') and not s.startswith('- ') and not s.startswith('#'):
            plus_num = 0

        # Main title: # ...
        if s.startswith('# ') and not s.startswith('## '):
            add_title(doc, s[2:])
            i += 1
            continue

        # Subtitle / version line (non-heading > line)
        if s.startswith('> ') and not s.startswith('>> '):
            txt = s[2:]
            if re.match(r'^(版本|Version|日期|Date)[：:]', txt, re.IGNORECASE):
                add_subtitle(doc, txt)
            elif txt.startswith('**') and txt.endswith('**'):
                add_subtitle(doc, txt)
            else:
                add_note(doc, txt)
            i += 1
            continue

        # Signature block: >> 发文机关, >> 日期
        if s.startswith('>> '):
            sig_lines = []
            while i < len(lines) and lines[i].rstrip('\n').strip().startswith('>> '):
                sig_lines.append(lines[i].rstrip('\n').strip()[3:])
                i += 1
            if sig_lines:
                add_signature(doc, sig_lines)
            continue

        # Separator
        if s == '---':
            i += 1
            continue

        # Chapter heading: ## 一、...
        if s.startswith('## '):
            h2_count = 0
            h3_count = 0
            heading_depth = 1
            bullet_num = 0
            plus_num = 0
            body_label_num = 0
            in_bold_section = False
            add_h1(doc, s[3:])
            i += 1
            continue

        # Section heading: ### ... → normalize to （一）（二）...
        if s.startswith('### '):
            h2_count += 1
            h3_count = 0
            heading_depth = 2
            bullet_num = 0
            plus_num = 0
            body_label_num = 0
            in_bold_section = False
            raw = s[4:]
            cleaned = _strip_heading_num(raw)
            add_h2(doc, f'（{_to_chinese_num(h2_count)}）{cleaned}')
            i += 1
            continue

        # Sub-section heading: #### ... → normalize to 1. 2. 3. ...
        if s.startswith('#### '):
            h3_count += 1
            heading_depth = 3
            bullet_num = 0
            plus_num = 0
            body_label_num = 0
            in_bold_section = False
            raw = s[5:]
            cleaned = _strip_heading_num(raw)
            add_h3(doc, f'{h3_count}.{cleaned}')
            i += 1
            continue

        # Code block (skip, render as plain body)
        if s.startswith('```'):
            i += 1
            while i < len(lines) and not lines[i].rstrip('\n').strip().startswith('```'):
                add_body(doc, lines[i].rstrip('\n'))
                i += 1
            i += 1
            continue

        # Table
        if '|' in s:
            headers, rows, i = _parse_table(lines, i)
            if headers and rows:
                add_table(doc, headers, rows, _is_content_table(rows))
            continue

        # Unordered list → context-aware numbering per GB/T 9704 hierarchy
        #   Outside bold section:
        #     depth 1-2 (under ## / ###): Level 3 → 1. 2. 3.
        #     depth 3   (under ####):     Level 4 → （1）（2）（3）
        #   Inside bold section (shift one level deeper):
        #     depth 1-2: Level 4 → （1）（2）（3）
        #     depth 3:   Level 5 → ①②③
        if s.startswith('- '):
            bullet_num += 1
            if in_bold_section:
                if heading_depth >= 3:
                    add_body(doc, f'{_circled_num(bullet_num)}{s[2:]}')
                else:
                    add_body(doc, f'（{bullet_num}）{s[2:]}')
            else:
                if heading_depth >= 3:
                    add_body(doc, f'（{bullet_num}）{s[2:]}')
                else:
                    add_body(doc, f'{bullet_num}.{s[2:]}')
            i += 1
            continue

        # Plus-list → Level 5: ①②③ (explicit circled-number items)
        if s.startswith('+ '):
            plus_num += 1
            add_body(doc, f'{_circled_num(plus_num)}{s[2:]}')
            i += 1
            continue

        # Ordered list: 1. ... or 1) ... → normalize spacing
        if re.match(r'^\d+[\.\)]\s', s):
            s = re.sub(r'^(\d+[\.\)])\s+', r'\1', s)
            add_body(doc, s)
            i += 1
            continue

        # Body text that starts with **NUMBER** → normalize prefix by depth
        # Bold labels act as virtual structural dividers:
        #   Under depth 1-2 (##/###): Level 3 → 1. 2. 3.
        #   Under depth 3   (####):   Level 4 → （1）（2）
        # Subsequent - items enter "bold section" and shift one level deeper.
        bold_prefix_m = re.match(r'^\*\*(.+?)\*\*', s)
        if bold_prefix_m:
            bold_part = bold_prefix_m.group(1)
            rest = s[bold_prefix_m.end():]
            cleaned = _strip_heading_num(bold_part)
            if cleaned != bold_part:
                body_label_num += 1
                if heading_depth >= 3:
                    prefix = f'（{body_label_num}）'
                else:
                    prefix = f'{body_label_num}.'
                s = f'**{prefix}{cleaned}**{rest}'
                in_bold_section = True
                bullet_num = 0
                plus_num = 0

        # Other body text
        add_body(doc, s)
        i += 1

    add_page_numbers(doc)

    doc.save(docx_path)
    return docx_path


# ── Auto-structure: weak-format → structured markdown ────────────

def _read_input_text(input_path, plain=False):
    """Read input file and return plain text, regardless of format.

    Supports .txt, .md (returned as-is), .docx (extracted via pandoc),
    .wps and .doc (cascading conversion: pandoc -> LibreOffice -> error).
    Set plain=True to strip ALL formatting (use for web-copied / messy docx).
    """
    ext = os.path.splitext(input_path)[1].lower()
    if ext == '.docx':
        import subprocess
        fmt = 'plain' if plain else 'markdown'
        result = subprocess.run(
            ['pandoc', input_path, '-t', fmt, '--wrap=none'],
            capture_output=True, text=True, encoding='utf-8',
            env={**os.environ, 'PYTHONIOENCODING': 'utf-8'})
        if result.returncode != 0:
            raise RuntimeError(f"pandoc failed: {result.stderr}")
        text = result.stdout
        # Undo pandoc's markdown escaping: \* → *, \[ → [, etc.
        text = text.replace('\\*', '*')
        text = text.replace('\\[', '[')
        text = text.replace('\\]', ']')
        text = text.replace('\\#', '#')
        return text
    elif ext in ('.wps', '.doc'):
        return _read_binary_doc(input_path, plain)
    else:
        with open(input_path, 'r', encoding='utf-8') as f:
            return f.read()


def _read_via_soffice(input_path, plain=False):
    """Convert .wps/.doc to text via LibreOffice headless conversion.

    Converts the binary document to a temporary .docx, then extracts text
    using pandoc (the same pipeline used for native .docx files).
    """
    import subprocess

    tmpdir = tempfile.mkdtemp(prefix='govdoc_')
    try:
        # Dynamically import the docx skill's soffice helper
        skills_dir = os.path.dirname(os.path.dirname(
            os.path.dirname(os.path.abspath(__file__))))
        soffice_py = os.path.join(
            skills_dir, 'docx', 'scripts', 'office', 'soffice.py')

        if os.path.exists(soffice_py):
            spec = importlib.util.spec_from_file_location(
                '_soffice_helper', soffice_py)
            mod = importlib.util.module_from_spec(spec)
            spec.loader.exec_module(mod)
            result = mod.run_soffice(
                ['--headless', '--convert-to', 'docx',
                 '--outdir', tmpdir, input_path],
                capture_output=True, text=True, timeout=120)
        else:
            result = subprocess.run(
                ['soffice', '--headless', '--convert-to', 'docx',
                 '--outdir', tmpdir, input_path],
                capture_output=True, text=True, timeout=120)

        if result.returncode != 0:
            raise RuntimeError(
                f"LibreOffice conversion failed (exit {result.returncode}): "
                f"{result.stderr.strip() or result.stdout.strip()}")

        base = os.path.splitext(os.path.basename(input_path))[0]
        tmp_docx = os.path.join(tmpdir, base + '.docx')
        if not os.path.exists(tmp_docx):
            raise RuntimeError(
                f"soffice did not produce expected output: {tmp_docx}")

        fmt = 'plain' if plain else 'markdown'
        result = subprocess.run(
            ['pandoc', tmp_docx, '-t', fmt, '--wrap=none'],
            capture_output=True, text=True, encoding='utf-8',
            env={**os.environ, 'PYTHONIOENCODING': 'utf-8'})
        if result.returncode != 0:
            raise RuntimeError(
                f"pandoc extraction failed after soffice conversion: "
                f"{result.stderr}")

        text = result.stdout
        text = text.replace('\\*', '*')
        text = text.replace('\\[', '[')
        text = text.replace('\\]', ']')
        text = text.replace('\\#', '#')
        return text

    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


def _read_binary_doc(input_path, plain=False):
    """Read .wps or .doc binary document via cascading conversion.

    Tries in order:
      1. pandoc (some installs support .doc/.wps via Word COM or custom readers)
      2. LibreOffice soffice (headless conversion to .docx, then pandoc)
      3. Helpful error message with installation instructions
    """
    import subprocess

    errors = []

    # Stage 1: try pandoc directly
    try:
        fmt = 'plain' if plain else 'markdown'
        result = subprocess.run(
            ['pandoc', input_path, '-t', fmt, '--wrap=none'],
            capture_output=True, text=True, encoding='utf-8',
            env={**os.environ, 'PYTHONIOENCODING': 'utf-8'})
        if result.returncode == 0 and result.stdout.strip():
            text = result.stdout
            text = text.replace('\\*', '*')
            text = text.replace('\\[', '[')
            text = text.replace('\\]', ']')
            text = text.replace('\\#', '#')
            return text
        else:
            errors.append(
                f"pandoc: {result.stderr.strip() or 'returned empty output'}")
    except FileNotFoundError:
        errors.append("pandoc: not found on PATH")
    except Exception as e:
        errors.append(f"pandoc: {e}")

    # Stage 2: try soffice (LibreOffice headless)
    try:
        return _read_via_soffice(input_path, plain)
    except FileNotFoundError:
        errors.append("soffice: LibreOffice not found on PATH")
    except Exception as e:
        errors.append(f"soffice: {e}")

    # Stage 3: all methods failed
    ext = os.path.splitext(input_path)[1]
    raise RuntimeError(
        f"Cannot read {ext} file '{os.path.basename(input_path)}'.\n"
        f"Conversion attempts:\n"
        + ''.join(f"  - {err}\n" for err in errors) +
        f"\nTo enable {ext} support, install LibreOffice:\n"
        f"  https://www.libreoffice.org/download/\n"
        f"Or manually convert {ext} to .docx first, then use --auto mode."
    )


def _is_body_num_pattern(text):
    """Check if a numbered line is body text (not a heading).

    Patterns like 一是/二是, 第一, 一方面, （1）（2） are body, not headings.
    """
    return bool(re.match(
        r'^[一二三四五六七八九十]+(是|方面|项|个|种|类|条|款|点|要)|'
        r'^第[一二三四五六七八九十]+[，,]|'
        r'^（\d+）|'
        r'^\(?\d+[\)）]|'
        r'^\d+(\.\d+)+',
        text))


def _bold_body_label(text):
    """Wrap ALL body-number labels (一是/二是/第一) in **bold** markers.

    Scans the entire text and wraps each label through its first
    period: "一是xxx。" → "**一是xxx。**"
    """
    label_re = (
        r'[一二三四五六七八九十]+(?:是|方面|项|个|种|类|条|款|点|要)'
        r'|第[一二三四五六七八九十]+[，,]'
    )
    full_re = re.compile(f'({label_re})')

    parts = []
    last_end = 0
    for m in full_re.finditer(text):
        label = m.group()
        start, end = m.start(), m.end()

        # Text before this label
        if start > last_end:
            parts.append(text[last_end:start])

        # Bold from label through first 。
        after = text[end:]
        period_idx = after.find('。')
        if period_idx >= 0 and period_idx <= 80:
            bold_part = label + after[:period_idx + 1]
            parts.append(f'**{bold_part}**')
            last_end = end + period_idx + 1
        else:
            # Just bold the label itself
            parts.append(f'**{label}**')
            last_end = end

    if last_end < len(text):
        parts.append(text[last_end:])

    return ''.join(parts) if parts else text


def _split_heading_body(text, max_heading=45):
    """If heading text is long, split at first 。into title + body."""
    if len(text) <= max_heading:
        return text, ''
    # Count 。in the text
    idx = text.find('。')
    if idx > 0 and idx < len(text) - 1:
        return text[:idx + 1], text[idx + 1:]
    return text, ''


def auto_structure(text):
    """Detect document structure from plain text and add markdown heading markers.

    Conventions detected:
      - First non-empty line (short, no numbering) → # main title
      - 一、二、三、… at line start  → ## (Level 1 chapter heading)
      - （一）（二）（三）… at line start → ### (Level 2 section heading)
      - N. … at line start after ###   → #### (Level 3 sub-section heading)
      - 一是/二是/第一/（1）/1.1.2 etc.  → kept as body text

    Returns markdown text ready for convert_markdown().
    """
    lines = text.split('\n')
    result = []
    in_level1 = False   # inside a ## section
    in_level2 = False   # inside a ### section
    title_done = False  # title already emitted
    has_chinese_headings = False  # seen any 一、 or （一）

    def _is_title_candidate(s):
        """First short non-numbered line is a title candidate."""
        return (len(s) >= 4 and len(s) <= 80
                and not re.match(r'^[一二三四五六七八九十（\d#]', s)
                and not s.startswith('>') and not s.startswith('|'))

    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            result.append('')
            continue

        # Already has markdown markers — strip pandoc ** and pass through
        if re.match(r'^#{1,4}\s', stripped):
            cleaned = stripped
            # Remove **bold** wrappers (pandoc artifact) — headings are bolded by their own handlers
            cleaned = re.sub(r'\*\*(.+?)\*\*', r'\1', cleaned)
            result.append(cleaned)
            title_done = True
            continue

        # Strip **bold** markers that pandoc wraps around headings
        # "**一、工作开展情况**" → "一、工作开展情况"
        plain = stripped
        bold_m = re.match(r'^\*\*(.+?)\*\*', stripped)
        if bold_m:
            after_bold = stripped[bold_m.end():]
            plain = bold_m.group(1) + after_bold

        # ── Detect main title: first non-empty line that looks like a title ──
        if not title_done and _is_title_candidate(plain):
            title_done = True
            result.append(f'# {plain}')
            result.append('')
            continue

        # ── Level 1: 一、二、三、… ──
        m1 = re.match(r'^([一二三四五六七八九十]+)、(.+)$', plain)
        if m1 and not _is_body_num_pattern(plain):
            title_done = True
            has_chinese_headings = True
            title, body = _split_heading_body(plain)
            in_level1 = True
            in_level2 = False
            result.append('')
            result.append(f'## {title}')
            if body:
                result.append('')
                result.append(_bold_body_label(body))
            continue

        # ── Level 2: （一）（二）（三）… ──
        m2 = re.match(r'^（([一二三四五六七八九十]+)）(.+)$', plain)
        if m2:
            title_done = True
            has_chinese_headings = True
            title, body = _split_heading_body(plain)
            in_level2 = True
            result.append(f'### {title}')
            if body:
                result.append('')
                result.append(_bold_body_label(body))
            continue

        # ── Level 3: N. … ──
        #   Inside ### section    → #### sub-section heading
        #   No Chinese headings   → ## chapter heading (arabic-only docs)
        #   Otherwise             → body text (pass through)
        m3 = re.match(r'^(\d+)\.\s*(.+)$', plain)
        if m3:
            if in_level2:
                result.append(f'#### {plain}')
            elif not in_level1 and not has_chinese_headings:
                title, body = _split_heading_body(plain)
                in_level1 = True
                result.append('')
                result.append(f'## {title}')
                if body:
                    result.append('')
                    result.append(_bold_body_label(body))
            else:
                result.append(stripped)
            continue

        # ── Regular body ──
        result.append(_bold_body_label(stripped))

    return '\n'.join(result)


def auto_convert(input_path, output_path):
    """Auto-detect structure and convert to formatted .docx.

    Reads .txt, .md, or .docx, auto-detects heading hierarchy,
    and converts to GB/T 9704-2012 formatted .docx.
    """
    text = _read_input_text(input_path)
    md_text = auto_structure(text)

    # Write temp markdown and convert
    tmp_md = os.path.join(os.path.dirname(output_path),
                          '_auto_temp.md')
    with open(tmp_md, 'w', encoding='utf-8') as f:
        f.write(md_text)

    try:
        result = convert_markdown(tmp_md, output_path)
    finally:
        if os.path.exists(tmp_md):
            os.remove(tmp_md)

    return result


# ── Legacy JSON mode ────────────────────────────────────────────

def generate_from_json(data):
    """Legacy JSON-input mode (kept for backward compatibility)."""
    doc = Document()

    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(3.7)
    sec.bottom_margin = Cm(3.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.6)

    # Default style
    ns = doc.styles['Normal']
    ns.font.name = FS
    ns.font.size = Pt(16)
    ns.element.rPr.rFonts.set(qn('w:eastAsia'), FS)
    ns.paragraph_format.line_spacing = Pt(29)
    ns.paragraph_format.space_before = Pt(0)
    ns.paragraph_format.space_after = Pt(0)
    ns.paragraph_format.first_line_indent = Pt(32)
    # Widow/orphan control
    ns._element.get_or_add_pPr().append(parse_xml(f'<w:widowControl {nsdecls("w")} w:val="0"/>'))

    if data.get('title'):
        add_title(doc, data['title'])
    if data.get('subtitle'):
        add_subtitle(doc, data['subtitle'])

    for line in data.get('content', '').split('\n'):
        s = line.strip()
        if not s:
            continue
        # Simple heuristic: detect heading levels by prefix
        if re.match(r'^[一二三四五六七八九十]+、', s):
            add_h1(doc, s)
        elif re.match(r'^（[一二三四五六七八九十]+）', s):
            add_h2(doc, s)
        elif s.startswith('> '):
            add_note(doc, s[2:])
        elif s.startswith('- '):
            add_bullet(doc, s[2:])
        else:
            add_body(doc, s)

    if data.get('page_numbers'):
        add_page_numbers(doc)

    doc.save(data['output_path'])
    return data['output_path']


# ── CLI ─────────────────────────────────────────────────────────

if __name__ == '__main__':
    parser = argparse.ArgumentParser(
        description='Convert markdown to Chinese government document (.docx)')
    parser.add_argument('--input', '-i', help='Path to input markdown file')
    parser.add_argument('--output', '-o', help='Path to output .docx file (optional)')
    parser.add_argument('--json_file', help='Path to JSON file (legacy mode)')
    parser.add_argument('--page-numbers', action='store_true',
                        help='Add page number footer')
    parser.add_argument('--auto', action='store_true',
                        help='Auto-detect structure from .txt / .docx / unformatted .md')
    args = parser.parse_args()

    if args.json_file:
        with open(args.json_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        out = generate_from_json(data)
        print(f'Saved: {out}')

    elif args.input:
        out = args.output
        if not out:
            base = os.path.splitext(args.input)[0]
            if args.auto:
                out = base + '_已排版.docx'
            else:
                out = base + '.docx'
        if args.auto:
            result = auto_convert(args.input, out)
        else:
            result = convert_markdown(args.input, out)
        print(f'Saved: {result}')

    else:
        parser.print_help()
